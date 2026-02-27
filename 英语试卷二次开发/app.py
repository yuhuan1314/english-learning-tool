# 英语试卷二次开发工具
# 运行命令: streamlit run app.py

import streamlit as st
import pandas as pd
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============ 页面配置 ============
st.set_page_config(
    page_title="英语试卷二次开发工具",
    page_icon="📝",
    layout="wide"
)

# ============ 样式 ============
st.markdown("""
<style>
    .main {
        background-color: #f5f5f5;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
    }
    .title {
        text-align: center;
        color: #2c5282;
        font-size: 32px;
        font-weight: bold;
    }
    .section-header {
        background-color: #4299e1;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
    .exercise-box {
        background-color: white;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #4299e1;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ============ 功能函数 ============

def extract_text_from_pdf(pdf_file):
    """从PDF提取文本"""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"PDF解析错误: {e}")
        return None

def identify_question_types(text):
    """识别题型"""
    types = []
    text_lower = text.lower()
    
    if "阅读理解" in text or "reading comprehension" in text_lower:
        types.append("阅读理解")
    if "七选五" in text or "七选五" in text:
        types.append("七选五")
    if "完形填空" in text or "cloze" in text_lower:
        types.append("完形填空")
    if "语法填空" in text or "grammar" in text_lower:
        types.append("语法填空")
    if "书面表达" in text or "writing" in text_lower:
        types.append("书面表达")
    if "听力" in text or "listening" in text_lower:
        types.append("听力")
    
    return types if types else ["阅读理解", "七选五", "完形填空", "语法填空"]

def extract_reading_articles(text):
    """提取阅读理解文章"""
    articles = []
    # 简单的文章提取逻辑
    lines = text.split('\n')
    current_article = []
    in_article = False
    
    for line in lines:
        if any(marker in line for marker in ['A\n', 'B\n', 'C\n', 'D\n']) and len(line) < 5:
            if current_article:
                articles.append('\n'.join(current_article))
                current_article = []
            in_article = True
        if in_article:
            current_article.append(line)
    
    if current_article:
        articles.append('\n'.join(current_article))
    
    return articles[:4]  # 最多4篇

def generate_vocabulary_exercise(article_text, topic):
    """生成词汇练习"""
    # 提取关键词
    words = re.findall(r'\b[a-zA-Z]{5,}\b', article_text.lower())
    word_freq = {}
    for w in words:
        word_freq[w] = word_freq.get(w, 0) + 1
    
    top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]
    
    return {
        "话题词汇": [w[0] for w in top_words],
        "动词": ["analyze", "explore", "discuss", "examine", "investigate"],
        "名词": ["concept", "principle", "theory", "perspective", "framework"],
        "形容词": ["significant", "essential", "fundamental", "crucial", "vital"]
    }

def generate_word_formation(article_text):
    """生成词性变形练习"""
    return {
        "动→名": ["discuss→discussion", "explain→explanation", "analyze→analysis"],
        "名→形": ["tradition→traditional", "culture→cultural", "nature→natural"],
        "副词": ["complete→completely", "true→truly", "happy→happily"]
    }

def generate_phrase_translation():
    """生成短语翻译"""
    return {
        "中文": ["不仅...而且...", "事实上", "重新审视", "纯粹的乐趣", "充实的生活"],
        "英文": ["not only...but also...", "in fact", "revisit the question of...", "for pure pleasure, not purpose", "a life well lived"]
    }

def generate_cloze_exercise(article_text):
    """生成语境填词"""
    words = ["however", "therefore", "although", "furthermore", "moreover", 
             "consequently", "nevertheless", "accordingly", "meanwhile", "otherwise"]
    selected = words[:8]
    
    exercise = {}
    for i, word in enumerate(selected, 1):
        exercise[f"{i}"] = {
            "word": word,
            "clue": f"选择适当的单词填空{i}"
        }
    return exercise

def generate_discourse_structure(article_text):
    """生成语篇结构"""
    return {
        "宏观骨架": [
            {"段落": "Para 1", "功能": "引言", "内容": "引出主题"},
            {"段落": "Para 2", "功能": "论点1", "内容": "主要论点"},
            {"段落": "Para 3", "功能": "论点2", "内容": "支持论据"},
            {"段落": "Para 4", "功能": "结论", "内容": "总结升华"}
        ],
        "衔接词": [
            {"原句": "...is important.", "挖空": "________, ... is important.", "答案": "Therefore", "逻辑": "因果"},
            {"原句": "However, ...", "挖空": "________, ...", "答案": "However", "逻辑": "转折"},
        ]
    }

def generate_micro_writing(topic, requirements):
    """生成微写作"""
    return {
        "情境": f"关于{topic}的写作",
        "词数": "50词左右",
        "要求": requirements,
        "开头提示": f"Write about {topic}...",
        "参考范文": f"This essay discusses {topic}. {topic} is an important concept that affects our daily life. In conclusion..."
    }

def generate_grammar_exercise(text):
    """生成语法练习"""
    return {
        "非谓语动词": ["现在分词作状语", "过去分词作定语", "不定式作目的状语"],
        "定语从句": ["which/that引导", "who/whom引导", "介词+which/whom"],
        "状语从句": ["时间状语", "原因状语", "让步状语", "结果状语"],
        "名词性从句": ["主语从句", "宾语从句", "表语从句", "同位语从句"]
    }

def create_word_document(exercises_data, article_title):
    """创建Word文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading(f'{article_title} - 二次开发练习', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 词汇练习
    doc.add_heading('一、话题词汇', level=1)
    if "vocabulary" in exercises_data:
        for category, words in exercises_data["vocabulary"].items():
            doc.add_paragraph(f"{category}: {', '.join(words)}")
    
    # 词性变形
    doc.add_heading('二、词性变形', level=1)
    if "word_formation" in exercises_data:
        for transform_type, examples in exercises_data["word_formation"].items():
            doc.add_paragraph(f"{transform_type}:")
            for ex in examples:
                doc.add_paragraph(f"  • {ex}", style='List Bullet')
    
    # 短语翻译
    doc.add_heading('三、短语翻译', level=1)
    if "phrase_translation" in exercises_data:
        table = doc.add_table(rows=len(exercises_data["phrase_translation"]["中文"])+1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = '中文'
        table.rows[0].cells[1].text = '英文'
        for i, (cn, en) in enumerate(zip(exercises_data["phrase_translation"]["中文"], 
                                          exercises_data["phrase_translation"]["英文"]), 1):
            table.rows[i].cells[0].text = cn
            table.rows[i].cells[1].text = en
    
    # 语篇结构
    doc.add_heading('四、语篇结构', level=1)
    doc.add_paragraph('宏观骨架：')
    if "discourse" in exercises_data and "宏观骨架" in exercises_data["discourse"]:
        for item in exercises_data["discourse"]["宏观骨架"]:
            doc.add_paragraph(f"  {item['段落']}: {item['功能']} - {item['内容']}")
    
    doc.add_paragraph('衔接词挖空：')
    if "discourse" in exercises_data and "衔接词" in exercises_data["discourse"]:
        for item in exercises_data["discourse"]["衔接词"]:
            doc.add_paragraph(f"  {item['挖空']} → 答案: {item['答案']} ({item['逻辑']})")
    
    # 微写作
    doc.add_heading('五、微写作', level=1)
    if "micro_writing" in exercises_data:
        mw = exercises_data["micro_writing"]
        doc.add_paragraph(f"情境: {mw['情境']}")
        doc.add_paragraph(f"词数: {mw['词数']}")
        doc.add_paragraph(f"要求: {mw['要求']}")
        doc.add_paragraph(f"开头提示: {mw['开头提示']}")
        doc.add_paragraph(f"参考范文: {mw['参考范文']}")
    
    # 保存
    return doc

# ============ 主界面 ============

st.markdown('<p class="title">📝 英语试卷二次开发工具</p>', unsafe_allow_html=True)

st.markdown("---")

# 侧边栏
with st.sidebar:
    st.header("⚙️ 设置")
    
    st.subheader("题型选择")
    question_types = st.multiselect(
        "选择要处理的题型",
        ["阅读理解", "七选五", "完形填空", "语法填空", "书面表达"],
        default=["阅读理解"]
    )
    
    st.sub_header("练习类型")
    include_vocab = st.checkbox("话题词汇", value=True)
    include_word_form = st.checkbox("词性变形", value=True)
    include_phrase = st.checkbox("短语翻译", value=True)
    include_cloze = st.checkbox("语境填词", value=True)
    include_discourse = st.checkbox("语篇结构", value=True)
    include_writing = st.checkbox("微写作", value=True)
    
    st.markdown("---")
    st.info("💡 提示：上传PDF试卷后，系统会自动识别题型并生成相应的二次开发练习。")

# 主区域
st.header("📤 上传试卷")

uploaded_file = st.file_uploader("选择PDF文件", type=['pdf'])

if uploaded_file:
    st.success("✅ 文件上传成功！")
    
    with st.spinner("🔍 正在解析试卷..."):
        text = extract_text_from_pdf(uploaded_file)
    
    if text:
        st.info(f"✅ 识别到试卷内容，长度: {len(text)} 字符")
        
        # 识别题型
        identified_types = identify_question_types(text)
        st.write(f"📋 识别到的题型: {', '.join(identified_types)}")
        
        # 显示选项
        selected_articles = st.selectbox(
            "选择文章进行二次开发",
            ["A篇", "B篇", "C篇", "D篇", "七选五", "完形填空", "语法填空"]
        )
        
        # 生成练习
        if st.button("🚀 生成二次开发练习", type="primary"):
            with st.spinner("✍️ 正在生成练习..."):
                # 构建练习数据
                exercises = {}
                
                if include_vocab:
                    exercises["vocabulary"] = generate_vocabulary_exercise(text, selected_articles)
                
                if include_word_form:
                    exercises["word_formation"] = generate_word_formation(text)
                
                if include_phrase:
                    exercises["phrase_translation"] = generate_phrase_translation()
                
                if include_cloze:
                    exercises["cloze"] = generate_cloze_exercise(text)
                
                if include_discourse:
                    exercises["discourse"] = generate_discourse_structure(text)
                
                if include_writing:
                    exercises["micro_writing"] = generate_micro_writing(
                        selected_articles, 
                        "运用本课话题词汇，使用With复合结构"
                    )
                
                # 显示练习
                st.markdown("---")
                st.header("📚 生成的练习")
                
                # 词汇
                if include_vocab and "vocabulary" in exercises:
                    st.markdown('<div class="section-header">一、话题词汇</div>', unsafe_allow_html=True)
                    for category, words in exercises["vocabulary"].items():
                        with st.expander(f"{category}"):
                            st.write(", ".join(words))
                
                # 词性变形
                if include_word_form and "word_formation" in exercises:
                    st.markdown('<div class="section-header">二、词性变形</div>', unsafe_allow_html=True)
                    for transform_type, examples in exercises["word_formation"].items():
                        st.markdown(f"**{transform_type}**")
                        for ex in examples:
                            st.write(f"  • {ex}")
                
                # 短语翻译
                if include_phrase and "phrase_translation" in exercises:
                    st.markdown('<div class="section-header">三、短语翻译</div>', unsafe_allow_html=True)
                    pt = exercises["phrase_translation"]
                    df = pd.DataFrame({"中文": pt["中文"], "英文": pt["英文"]})
                    st.table(df)
                
                # 语篇结构
                if include_discourse and "discourse" in exercises:
                    st.markdown('<div class="section-header">四、语篇结构</div>', unsafe_allow_html=True)
                    
                    st.subheader("宏观骨架")
                    ds = exercises["discourse"]
                    if "宏观骨架" in ds:
                        for item in ds["宏观骨架"]:
                            st.write(f"**{item['段落']}**: {item['功能']} - {item['内容']}")
                    
                    st.subheader("衔接词挖空")
                    if "衔接词" in ds:
                        for item in ds["衔接词"]:
                            st.text_input(
                                item['挖空'], 
                                value=item['答案'],
                                help=f"逻辑: {item['逻辑']}"
                            )
                
                # 微写作
                if include_writing and "micro_writing" in exercises:
                    st.markdown('<div class="section-header">五、微写作</div>', unsafe_allow_html=True)
                    mw = exercises["micro_writing"]
                    st.write(f"**情境**: {mw['情境']}")
                    st.write(f"**词数**: {mw['词数']}")
                    st.write(f"**要求**: {mw['要求']}")
                    st.write(f"**开头提示**: {mw['开头提示']}")
                    with st.expander("参考范文"):
                        st.write(mw['参考范文'])
                
                # 导出
                st.markdown("---")
                st.header("💾 导出")
                
                # 创建Word文档
                doc = create_word_document(exercises, selected_articles)
                
                # 保存到内存
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                st.download_button(
                    label="📥 下载Word文档",
                    data=doc_buffer,
                    file_name=f"{selected_articles}_二次开发练习.docx",
                    type="primary"
                )

else:
    st.info("👆 请上传PDF格式的英语试卷文件")
    
    # 示例
    st.markdown("---")
    st.header("📋 示例")
    st.write("""
    ### 生成的练习包含：
    
    1. **话题词汇** - 动作引擎、语义支柱、修饰词
    2. **词性变形** - 动→名、名→形、副词
    3. **短语翻译** - 整块逻辑短语
    4. **语境填词** - 基于文章的挖空练习
    5. **语篇结构** - 宏观骨架、衔接词挖空
    6. **微写作** - 50词情境写作
    """)

# 页脚
st.markdown("---")
st.markdown("<div style='text-align: center; color: #666;'>英语试卷二次开发工具 | 基于克拉申输入假说与布鲁姆认知模型</div>", unsafe_allow_html=True)
